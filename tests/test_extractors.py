"""格式抽取测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from legal_regulation_slicer.extractors import (
    DocPreconversionRequiredError,
    discover_candidate_files,
    extract_text_from_file,
)


def test_extract_text_from_docx_reads_paragraphs_and_table_cells(tmp_path: Path) -> None:
    """docx 应提取段落与表格单元格文本。"""

    file_path = tmp_path / "法规示例.docx"
    document = Document()
    document.add_paragraph("第一条 为了规范抽取。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "第二条 表格中的法条内容。"
    document.save(str(file_path))

    text = extract_text_from_file(file_path)

    assert "第一条 为了规范抽取。" in text
    assert "第二条 表格中的法条内容。" in text


def test_extract_text_from_txt_and_md_reads_common_text_files(tmp_path: Path) -> None:
    """txt 与 md 应支持常见中文文本编码。"""

    txt_path = tmp_path / "法规示例.txt"
    md_path = tmp_path / "法规示例.md"
    txt_path.write_bytes("第一条 文本法规内容。".encode("gb18030"))
    md_path.write_text("# 示例\n第二条 Markdown 法规内容。", encoding="utf-8")

    txt_text = extract_text_from_file(txt_path)
    md_text = extract_text_from_file(md_path)

    assert "第一条 文本法规内容。" in txt_text
    assert "第二条 Markdown 法规内容。" in md_text


def test_extract_text_from_pdf_uses_pdfplumber_pages(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    """pdf 应通过 pdfplumber 页面抽取文本。"""

    class _FakePage:
        """模拟 pdf 页面对象。"""

        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            """返回页面文本。"""

            return self._text

    class _FakePdf:
        """模拟 pdfplumber 打开的上下文对象。"""

        def __init__(self) -> None:
            self.pages = [_FakePage("第一条 PDF 法规内容。"), _FakePage("第二条 PDF 第二页。")]

        def __enter__(self) -> "_FakePdf":
            """进入上下文。"""

            return self

        def __exit__(self, exc_type, exc, tb) -> bool:
            """退出上下文。"""

            return False

    def _fake_open(_file_path: Path) -> _FakePdf:
        """替换 pdfplumber.open。"""

        return _FakePdf()

    monkeypatch.setattr("legal_regulation_slicer.extractors.pdfplumber.open", _fake_open)
    pdf_path = tmp_path / "法规示例.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")

    text = extract_text_from_file(pdf_path)

    assert "第一条 PDF 法规内容。" in text
    assert "第二条 PDF 第二页。" in text


def test_extract_text_from_doc_requires_preconversion(tmp_path: Path) -> None:
    """doc 应明确提示先转为 docx。"""

    file_path = tmp_path / "旧版法规.doc"
    file_path.write_bytes(b"legacy-doc")

    with pytest.raises(DocPreconversionRequiredError) as exc_info:
        extract_text_from_file(file_path)

    assert ".doc" in str(exc_info.value)
    assert "docx" in str(exc_info.value).lower()


def test_discover_candidate_files_skips_excluded_decision_directories(tmp_path: Path) -> None:
    """“修改、废止的决定”目录应在候选发现阶段被整棵跳过。"""

    input_root = tmp_path / "法律法规数据库"
    included = input_root / "法律" / "法律" / "刑法" / "中华人民共和国刑法_20201226.docx"
    excluded = input_root / "法律" / "修改、废止的决定" / "关于修改某法的决定_20200101.docx"
    included.parent.mkdir(parents=True)
    excluded.parent.mkdir(parents=True)

    document = Document()
    document.add_paragraph("第一条 为了规范抽取。")
    document.save(str(included))
    document.save(str(excluded))

    discovered = discover_candidate_files(input_root, recursive=True)

    assert discovered == [included.resolve()]
