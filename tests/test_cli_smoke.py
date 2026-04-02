"""CLI 冒烟测试。"""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document


def test_cli_outputs_articles_manifest_and_failed_files(tmp_path: Path) -> None:
    """批量执行后应生成法条、清单与失败文件三类产物。"""

    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()

    docx_path = input_dir / "法律" / "法律" / "刑法" / "反电信网络诈骗法_20221201.docx"
    docx_path.parent.mkdir(parents=True)
    document = Document()
    document.add_paragraph("目录")
    document.add_paragraph("第一章 总则")
    document.add_paragraph("第一条 为了预防电信网络诈骗。")
    document.add_paragraph("第二条之一 本法适用于有关活动。")
    document.save(str(docx_path))

    md_path = input_dir / "行政法规" / "行政法规" / "治安管理处罚法_20060301.md"
    md_path.parent.mkdir(parents=True)
    md_path.write_text(
        "第一条 为维护社会治安秩序。\n第二条 本法适用于扰乱公共秩序等行为。",
        encoding="utf-8",
    )

    legacy_doc_path = input_dir / "法律" / "法律解释" / "旧版法规_20200101.doc"
    legacy_doc_path.parent.mkdir(parents=True)
    legacy_doc_path.write_bytes(b"legacy-doc")

    excluded_docx_path = input_dir / "法律" / "修改、废止的决定" / "关于修改某法的决定_20210101.docx"
    excluded_docx_path.parent.mkdir(parents=True)
    excluded_document = Document()
    excluded_document.add_paragraph("第一条 本决定略。")
    excluded_document.save(str(excluded_docx_path))

    script_path = Path(__file__).resolve().parents[1] / "scripts" / "slice_regulations.py"
    completed = subprocess.run(
        [
            sys.executable,
            str(script_path),
            "--input",
            str(input_dir),
            "--output",
            str(output_dir),
            "--recursive",
        ],
        capture_output=True,
        text=True,
    )

    assert completed.returncode == 0, completed.stderr

    articles_path = output_dir / "articles.jsonl"
    manifest_path = output_dir / "manifest.json"
    failed_path = output_dir / "failed_files.jsonl"
    assert articles_path.exists()
    assert manifest_path.exists()
    assert failed_path.exists()

    articles = [
        json.loads(line)
        for line in articles_path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    failed = [
        json.loads(line)
        for line in failed_path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]

    assert manifest["candidate_files"] == 3
    assert manifest["processed_files"] == 2
    assert manifest["failed_files"] == 1
    assert manifest["article_count"] == 4
    assert len(articles) == 4
    assert failed[0]["reason_code"] == "doc_requires_preconversion"
    assert failed[0]["source_file"].endswith("法律/法律解释/旧版法规_20200101.doc")
    assert {item["law_name"] for item in articles} == {"反电信网络诈骗法", "治安管理处罚法"}
    assert all("law_domain" not in item for item in articles)
    assert all("source_category" not in item for item in articles)
    assert all("source_region" not in item for item in articles)
    assert {item["legal_category"] for item in articles} == {"法律", "行政法规"}
    assert {item["legal_subcategory"] for item in articles} == {"法律", "行政法规"}
    assert {item["legal_topic"] for item in articles if item["law_name"] == "反电信网络诈骗法"} == {"刑法"}
    assert {item["legal_topic"] for item in articles if item["law_name"] == "治安管理处罚法"} == {None}
    assert set(manifest["cli_args"]) == {"legal_category", "legal_subcategory", "legal_topic", "authority_default"}
