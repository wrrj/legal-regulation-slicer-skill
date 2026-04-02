"""文件发现与文本抽取。"""

from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document

from .models import (
    DocPreconversionRequiredError,
    EmptyExtractionError,
    ExtractionError,
    UnsupportedFileTypeError,
)

SUPPORTED_FILE_SUFFIXES = {".docx", ".pdf", ".txt", ".md"}
DISCOVERABLE_FILE_SUFFIXES = SUPPORTED_FILE_SUFFIXES | {".doc"}
TEXT_FILE_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "gbk")
EXCLUDED_DIRECTORY_NAMES = {"修改、废止的决定"}


def discover_candidate_files(input_path: Path, recursive: bool = False) -> list[Path]:
    """发现可处理候选文件。"""

    resolved_input = input_path.resolve()
    if resolved_input.is_file():
        if any(parent.name in EXCLUDED_DIRECTORY_NAMES for parent in resolved_input.parents):
            return []
        return [resolved_input]

    pattern = "**/*" if recursive else "*"
    candidates = [
        file_path.resolve()
        for file_path in resolved_input.glob(pattern)
        if file_path.is_file()
        and file_path.suffix.lower() in DISCOVERABLE_FILE_SUFFIXES
        and not _is_in_excluded_directory(file_path, resolved_input)
    ]
    candidates.sort(key=lambda item: str(item))
    return candidates


def _is_in_excluded_directory(file_path: Path, input_root: Path) -> bool:
    """判断文件是否位于应被跳过的目录子树中。"""

    relative_parts = []
    try:
        relative_parts = file_path.resolve().relative_to(input_root.resolve()).parts[:-1]
    except ValueError:
        relative_parts = [part.name for part in file_path.parents]
    return any(part in EXCLUDED_DIRECTORY_NAMES for part in relative_parts)


def extract_text_from_file(file_path: Path) -> str:
    """根据文件后缀提取纯文本。"""

    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        text = _extract_text_from_docx(file_path)
    elif suffix == ".pdf":
        text = _extract_text_from_pdf(file_path)
    elif suffix in {".txt", ".md"}:
        text = _extract_text_from_plain_text(file_path)
    elif suffix == ".doc":
        raise DocPreconversionRequiredError(file_path.name)
    else:
        raise UnsupportedFileTypeError(suffix)

    cleaned = text.strip()
    if not cleaned:
        raise EmptyExtractionError(file_path.name)
    return cleaned


def _extract_text_from_docx(file_path: Path) -> str:
    """抽取 docx 段落与表格单元格文本。"""

    document = Document(str(file_path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line:
            blocks.append(line)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                line = cell.text.strip()
                if line:
                    blocks.append(line)
    return "\n".join(blocks)


def _extract_text_from_pdf(file_path: Path) -> str:
    """抽取文本型 PDF 的页面文本。"""

    blocks: list[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            line = (page.extract_text() or "").strip()
            if line:
                blocks.append(line)
    return "\n".join(blocks)


def _extract_text_from_plain_text(file_path: Path) -> str:
    """按常见中文编码读取纯文本文件。"""

    last_error: Exception | None = None
    for encoding in TEXT_FILE_ENCODINGS:
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
            continue
    raise ExtractionError(
        "text_decode_failed",
        f"文件 {file_path.name} 使用常见编码均无法解码：{last_error}",
    )
